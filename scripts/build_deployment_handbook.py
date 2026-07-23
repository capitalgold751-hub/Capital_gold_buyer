from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FIREBASE_VERCEL_GODADDY_DEPLOYMENT_GUIDE.md"
OUTPUT = ROOT / "docs" / "Capital_Gold_Buyers_Firebase_Vercel_Deployment_Handbook.docx"
LOGO = ROOT / "public" / "icons" / "icon-512.png"

GREEN = "075143"
DARK = "143A32"
GOLD = "D5A529"
PALE_GOLD = "F8F1DB"
PALE_GREEN = "EAF3F0"
MUTED = RGBColor(76, 94, 88)


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Capital Gold Buyers  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor(30, 48, 43)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (("Title", 32, DARK), ("Heading 1", 19, GREEN), ("Heading 2", 13.5, DARK), ("Heading 3", 11, GREEN)):
        style = document.styles[style_name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "CAPITAL GOLD BUYERS   /   FIREBASE OPERATIONS"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(7.5)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(GREEN)
    add_page_number(section.footer.paragraphs[0])


def clean_inline(text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("__", "")
    return text


def add_inline(paragraph, text):
    text = clean_inline(text)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.3)
            run.font.color.rgb = RGBColor.from_string(GREEN)
        else:
            paragraph.add_run(part)


def add_code_block(document, lines):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, "F3F5F4")
    set_cell_margins(cell, 120, 140, 120, 140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(7.7)
        run.font.color.rgb = RGBColor(31, 51, 45)
        if index < len(lines) - 1:
            run.add_break()
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table_row(line):
    return [clean_inline(item.strip()).replace("`", "") for item in line.strip().strip("|").split("|")]


def add_table(document, raw_lines):
    rows = [parse_table_row(line) for line in raw_lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", value.replace(" ", "")) for value in rows[1]):
        rows.pop(1)
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            shade(cell, GREEN if row_index == 0 else ("F7FAF9" if row_index % 2 == 0 else "FFFFFF"))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(values[column_index] if column_index < len(values) else "")
            run.font.name = "Aptos"
            run.font.size = Pt(7.7)
            run.font.bold = row_index == 0
            run.font.color.rgb = RGBColor(255, 255, 255) if row_index == 0 else RGBColor(34, 53, 47)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(20)
    paragraph.add_run().add_picture(str(LOGO), width=Inches(1.55))

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("PRODUCTION DELIVERY HANDBOOK")
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GOLD)
    run.font.letter_spacing = Pt(1.2)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Firebase + Vercel\nDeployment Handbook")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(8)
    run = subtitle.add_run("Authentication  •  Firestore  •  GoldAPI  •  Email  •  GoDaddy  •  PWA")
    run.font.name = "Georgia"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    shade(cell, PALE_GOLD)
    set_cell_margins(cell, 220, 260, 220, 260)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Prepared for capitalgoldbuyers.in\n").bold = True
    detail = paragraph.add_run("Version 2.0  •  Firebase migration edition  •  19 July 2026")
    detail.font.size = Pt(9)
    detail.font.color.rgb = MUTED

    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_before = Pt(20)
    run = warning.add_run("CONFIDENTIAL OPERATIONS GUIDE")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    document.add_page_break()


def add_contents(document, markdown):
    document.add_heading("Contents", level=1)
    for line in markdown.splitlines():
        if line.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Normal"]
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.space_after = Pt(2.2)
            run = paragraph.add_run(clean_inline(line[3:]))
            run.font.size = Pt(8.8)
            run.font.color.rgb = RGBColor.from_string(GREEN)
    document.add_page_break()


def add_markdown(document, markdown):
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_table(document, table_lines)
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(clean_inline(stripped[3:]).replace("`", ""), level=1)
        elif stripped.startswith("### "):
            document.add_heading(clean_inline(stripped[4:]).replace("`", ""), level=2)
        elif stripped.startswith("#### "):
            document.add_heading(clean_inline(stripped[5:]).replace("`", ""), level=3)
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.19)
            add_inline(paragraph, stripped)
        elif stripped.startswith("- [ ] "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.first_line_indent = Inches(-0.14)
            add_inline(paragraph, "☐ " + stripped[6:])
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
        else:
            paragraph = document.add_paragraph()
            add_inline(paragraph, stripped)
        index += 1


def build():
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    add_cover(document)
    add_contents(document, markdown)
    add_markdown(document, markdown)
    document.core_properties.title = "Capital Gold Buyers Firebase + Vercel Deployment Handbook"
    document.core_properties.subject = "Production Firebase, Firestore, authentication, Vercel and GoDaddy setup"
    document.core_properties.author = "Capital Gold Buyers"
    document.core_properties.keywords = "Firebase, Firestore, Vercel, GoDaddy, PWA, GoldAPI, RBAC"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
